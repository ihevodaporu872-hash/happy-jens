"""
Export Client for Telegram Bot

Exports bot responses to PDF and DOCX formats.
"""

import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


class ExportClient:
    """
    Client for exporting text content to PDF and DOCX formats.
    """

    def __init__(self, temp_dir: Optional[Path] = None):
        """
        Initialize the export client.

        Args:
            temp_dir: Directory for temporary files (default: system temp)
        """
        if temp_dir:
            self.temp_dir = temp_dir
        else:
            import tempfile
            self.temp_dir = Path(tempfile.gettempdir()) / "bot_exports"
        self.temp_dir.mkdir(parents=True, exist_ok=True)

    def _clean_markdown(self, text: str) -> str:
        """Remove markdown formatting for plain text export."""
        # Remove bold
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        # Remove italic
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        # Remove code blocks
        text = re.sub(r'```.*?```', '', text, flags=re.DOTALL)
        # Remove inline code
        text = re.sub(r'`(.*?)`', r'\1', text)
        return text

    def _generate_filename(self, title: str, extension: str) -> str:
        """Generate unique filename."""
        # Clean title for filename
        clean_title = re.sub(r'[^\w\s-]', '', title)[:30]
        clean_title = re.sub(r'\s+', '_', clean_title)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"{clean_title}_{timestamp}{extension}"

    def export_to_pdf(
        self,
        content: str,
        title: str = "Export",
        question: str = "",
        store_name: str = ""
    ) -> Optional[Path]:
        """
        Export content to PDF file.

        Args:
            content: Text content to export
            title: Document title
            question: Original question (optional)
            store_name: Source store name (optional)

        Returns:
            Path to created PDF file or None on failure
        """
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

            filename = self._generate_filename(title, ".pdf")
            filepath = self.temp_dir / filename

            doc = SimpleDocTemplate(
                str(filepath),
                pagesize=A4,
                rightMargin=2*cm,
                leftMargin=2*cm,
                topMargin=2*cm,
                bottomMargin=2*cm
            )

            styles = getSampleStyleSheet()

            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=12,
                textColor=colors.HexColor('#1a1a2e')
            )

            meta_style = ParagraphStyle(
                'Meta',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#666666'),
                spaceAfter=6
            )

            body_style = ParagraphStyle(
                'Body',
                parent=styles['Normal'],
                fontSize=11,
                leading=16,
                spaceAfter=8
            )

            story = []

            # Title
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 0.3*cm))

            # Metadata
            timestamp = datetime.now().strftime("%d.%m.%Y %H:%M")
            story.append(Paragraph(f"Дата: {timestamp}", meta_style))

            if store_name:
                story.append(Paragraph(f"Источник: {store_name}", meta_style))

            story.append(Spacer(1, 0.5*cm))

            # Question if provided
            if question:
                story.append(Paragraph(f"<b>Вопрос:</b> {question}", body_style))
                story.append(Spacer(1, 0.3*cm))

            # Content - split by paragraphs
            clean_content = self._clean_markdown(content)
            paragraphs = clean_content.split('\n\n')

            for para in paragraphs:
                para = para.strip()
                if para:
                    # Handle headers (lines starting with ##)
                    if para.startswith('#'):
                        para = re.sub(r'^#+\s*', '', para)
                        story.append(Paragraph(f"<b>{para}</b>", body_style))
                    else:
                        # Replace single newlines with <br/>
                        para = para.replace('\n', '<br/>')
                        story.append(Paragraph(para, body_style))
                    story.append(Spacer(1, 0.2*cm))

            # Footer
            story.append(Spacer(1, 1*cm))
            footer_text = "Generated by Gemini 3 Flash Bot"
            story.append(Paragraph(footer_text, meta_style))

            doc.build(story)
            logger.info(f"Exported PDF: {filepath}")
            return filepath

        except ImportError:
            logger.error("reportlab not installed. Install with: pip install reportlab")
            return None
        except Exception as e:
            logger.error(f"PDF export failed: {e}")
            return None

    def export_to_docx(
        self,
        content: str,
        title: str = "Export",
        question: str = "",
        store_name: str = ""
    ) -> Optional[Path]:
        """
        Export content to DOCX file.

        Args:
            content: Text content to export
            title: Document title
            question: Original question (optional)
            store_name: Source store name (optional)

        Returns:
            Path to created DOCX file or None on failure
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            filename = self._generate_filename(title, ".docx")
            filepath = self.temp_dir / filename

            doc = Document()

            # Title
            title_para = doc.add_heading(title, 0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Metadata
            timestamp = datetime.now().strftime("%d.%m.%Y %H:%M")
            meta = doc.add_paragraph()
            meta.add_run(f"Дата: {timestamp}").italic = True

            if store_name:
                meta.add_run(f"\nИсточник: {store_name}").italic = True

            doc.add_paragraph()  # Spacer

            # Question if provided
            if question:
                q_para = doc.add_paragraph()
                q_para.add_run("Вопрос: ").bold = True
                q_para.add_run(question)
                doc.add_paragraph()

            # Content
            clean_content = self._clean_markdown(content)
            paragraphs = clean_content.split('\n\n')

            for para in paragraphs:
                para = para.strip()
                if para:
                    # Handle headers
                    if para.startswith('#'):
                        level = len(re.match(r'^#+', para).group())
                        para = re.sub(r'^#+\s*', '', para)
                        doc.add_heading(para, level=min(level, 3))
                    else:
                        p = doc.add_paragraph()
                        # Handle bold text within paragraph
                        parts = re.split(r'(\*\*.*?\*\*)', para)
                        for part in parts:
                            if part.startswith('**') and part.endswith('**'):
                                p.add_run(part[2:-2]).bold = True
                            else:
                                p.add_run(part)

            # Footer
            doc.add_paragraph()
            footer = doc.add_paragraph()
            footer.add_run("Generated by Gemini 3 Flash Bot").italic = True
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.save(str(filepath))
            logger.info(f"Exported DOCX: {filepath}")
            return filepath

        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            return None
        except Exception as e:
            logger.error(f"DOCX export failed: {e}")
            return None

    def cleanup_old_files(self, hours: int = 24):
        """Remove export files older than specified hours."""
        from datetime import timedelta
        cutoff = datetime.now() - timedelta(hours=hours)

        cleaned = 0
        for file in self.temp_dir.glob("*"):
            if file.is_file():
                mtime = datetime.fromtimestamp(file.stat().st_mtime)
                if mtime < cutoff:
                    file.unlink()
                    cleaned += 1

        if cleaned > 0:
            logger.info(f"Cleaned up {cleaned} old export files")
